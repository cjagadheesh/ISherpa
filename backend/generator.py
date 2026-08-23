import logging
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# Configure logger
logger = logging.getLogger("sebi-ipo-generator.generator")

# Custom XML styling helpers for python-docx (shading, padding, borders)
def set_cell_background(cell, hex_color: str):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=80, bottom=80, left=100, right=100):
    """Sets internal padding (margins) of a cell in dxa (1/20 of a pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def _is_empty_value(val: Any) -> bool:
    """Mirrors validator._is_empty_value — list/table fields are empty only when they hold zero items."""
    if val is None:
        return True
    if isinstance(val, bool):
        return False
    if isinstance(val, (list, dict)):
        return len(val) == 0
    return str(val).strip() == ""


# Exact color constants matching the prospectus layout:
#   D09E73 tan      — Banners & Cover Page headers
#   D9D9D9 lt. gray — Column headers of summary content tables
TABLE_HEADER_TAN = "D09E73"
CONTENT_HEADER_GRAY = "D9D9D9"


def _set_horizontal_rule_borders(table):
    """Applies a single 0.5pt black horizontal rule on the TOP and BOTTOM of every cell only
    (no vertical lines), matching the exact prospectus print layout."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')
            for edge in ('top', 'bottom'):
                el = OxmlElement(f'w:{edge}')
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), '000000')
                borders.append(el)
            tcPr.append(borders)


def _add_grid_table(doc, headers: list, data_rows: list, header_bg: str = CONTENT_HEADER_GRAY):
    """Generates a table with formatted headers and horizontal-rule-only borders."""
    table = doc.add_table(rows=len(data_rows) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for c, h in enumerate(headers):
        header_cells[c].text = str(h)
        if header_cells[c].paragraphs[0].runs:
            header_cells[c].paragraphs[0].runs[0].font.bold = True
            header_cells[c].paragraphs[0].runs[0].font.size = Pt(9.5)
            header_cells[c].paragraphs[0].runs[0].font.color.rgb = RGBColor(20, 20, 20)
        set_cell_background(header_cells[c], header_bg)
        set_cell_margins(header_cells[c], 80, 80, 100, 100)
    for r, row in enumerate(data_rows):
        cells = table.rows[r + 1].cells
        for c, val in enumerate(row):
            cells[c].text = "" if val is None else str(val)
            if cells[c].paragraphs[0].runs:
                cells[c].paragraphs[0].runs[0].font.size = Pt(9.5)
            set_cell_margins(cells[c], 60, 60, 100, 100)
    _set_horizontal_rule_borders(table)
    return table


def _add_banner(doc, title: str):
    """Full-width tan banner heading used for major block headings."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    cell.text = title.upper()
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p.runs:
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(10.5)
        p.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    set_cell_background(cell, TABLE_HEADER_TAN)
    set_cell_margins(cell, 60, 60, 100, 100)
    _set_horizontal_rule_borders(table)
    return table


def _missing_marker(doc, text: str):
    warn = doc.add_paragraph()
    rw = warn.add_run(text)
    rw.font.color.rgb = RGBColor(220, 38, 38)
    rw.font.italic = True
    return warn


def _clean_llm_markdown(text: str) -> str:
    if not isinstance(text, str):
        return text
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("#"):
            stripped = stripped.lstrip("#").strip()
            if not stripped:
                continue
        stripped = stripped.replace("**", "")
        lines.append(stripped)
    return "\n".join(lines).strip()


def _add_multi_year_table(doc, metric_defs: list, merged: Dict[str, Any]):
    fy_order = []
    for _, key in metric_defs:
        rows = merged.get(key)
        if isinstance(rows, list):
            for row in rows:
                if isinstance(row, dict) and row.get("fy") and row["fy"] not in fy_order:
                    fy_order.append(row["fy"])
    if not fy_order:
        fy_order = ["Fiscal 2026", "Fiscal 2025", "Fiscal 2024"]
    
    data_rows = []
    for label, key in metric_defs:
        rows = merged.get(key)
        by_fy = {row.get("fy"): row.get("value") for row in rows if isinstance(row, dict)} if isinstance(rows, list) else {}
        data_rows.append([label] + [str(by_fy.get(fy, "[●]")) for fy in fy_order])
    _add_grid_table(doc, ["Particulars"] + fy_order, data_rows)


def _heading_numbered(doc, number, title: str):
    p = doc.add_paragraph()
    r = p.add_run(f"{number}. {title}")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _heading_lettered(doc, letter: str, title: str):
    p = doc.add_paragraph()
    r = p.add_run(f"{letter}) {title}")
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _add_notes_block(doc, notes: list):
    """Renders a numbered 'Notes:' block in the small italic style every draft/*.docx sample
    uses under its Financial Information and KPI tables — standard SEBI-prescribed ratio/
    definition boilerplate that applies to any company's restated financials, not just the
    one currently loaded, so it's generated here rather than requiring wizard input."""
    label_p = doc.add_paragraph()
    r_label = label_p.add_run("Notes:")
    r_label.font.bold = True
    r_label.font.size = Pt(9)
    for i, note in enumerate(notes, start=1):
        p = doc.add_paragraph(f"{i}. {note}")
        for run in p.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True


# Standard SEBI-prescribed KPI definitions, keyed by KPI name exactly as they appear in
# schema.json's kpi_sector_templates — only definitions matching KPIs actually present in
# kpi_values get rendered, so a Manufacturing filing doesn't show NBFC-only definitions.
KPI_DEFINITIONS = {
    "Revenue from Operations": "Revenue from operations is as per the Restated Financial Statements.",
    "EBITDA": "EBITDA is calculated as profit before tax plus finance costs, depreciation and amortization expense, less other income.",
    "EBITDA Margin": "EBITDA Margin (%) is calculated as EBITDA divided by revenue from operations.",
    "PAT": "PAT is profit for the year as per the Restated Financial Statements.",
    "PAT Margin": "PAT Margin (%) is calculated as profit for the year divided by revenue from operations.",
    "Return on Equity": "Return on Equity (%) is calculated as profit for the year divided by average total equity, where average total equity is the simple average of opening and closing total equity for the year.",
    "Return on Capital Employed": "Return on Capital Employed (%) is calculated as earnings before interest and tax (EBIT) divided by average capital employed. Capital employed is calculated as total assets less current liabilities.",
    "Trade Receivable Days": "Trade Receivable Days is calculated as average trade receivables divided by revenue from operations, multiplied by 365.",
    "Debtor Days": "Debtor Days is calculated as average trade receivables divided by revenue from operations, multiplied by 365.",
    "Trade Payable Days": "Trade Payable Days is calculated as average trade payables divided by cost of goods sold, multiplied by 365.",
    "Creditor Days": "Creditor Days is calculated as average trade payables divided by cost of goods sold, multiplied by 365.",
    "Inventory Days": "Inventory Days is calculated as average inventories divided by cost of goods sold, multiplied by 365.",
    "Working Capital Cycle": "Working Capital Cycle is calculated as Debtor/Receivable Days plus Inventory Days, less Creditor/Payable Days.",
    "Cash Conversion Cycle": "Cash Conversion Cycle is calculated as Inventory Days plus Trade Receivable Days, less Trade Payable Days.",
    "Inventory Turnover Ratio": "Inventory Turnover Ratio is calculated as cost of goods sold divided by average inventories.",
    "Sales to Retained Customers": "Sales to Retained Customers represents revenue from operations generated during the relevant period from customers continuously associated with the Company for three years or more.",
    "Ratio of Sales through Retained Customers": "Ratio of Sales through Retained Customers (%) is calculated as sales to retained customers divided by revenue from operations.",
    "Sales Volume": "Sales Volume represents the aggregate quantity sold by the Company during the relevant period.",
    "Debt to Equity": "Debt to Equity is calculated as Total Borrowings divided by total equity.",
    "Net Debt to EBITDA": "Net Debt to EBITDA is calculated as Net Debt (Total Borrowings less cash and cash equivalents) divided by EBITDA.",
    "Order Book": "Order Book is the total value of confirmed orders received but not yet executed or delivered as of the relevant date.",
    "Loans (AUM)": "Loans (AUM) represents the aggregate of Loans and Impairment loss allowance on loans as of the last day of the relevant year.",
    "Loans (AUM) Growth": "Loans (AUM) Growth represents the percentage growth in Loans (AUM) over the preceding year.",
    "Disbursements": "Disbursements represent the aggregate of all loan amounts extended to customers during the relevant year.",
    "Yield on Average Loans (AUM)": "Yield on Average Loans (AUM) represents interest income on loans for the relevant year as a percentage of Average Loans (AUM).",
    "Average Cost of Borrowings": "Average Cost of Borrowings represents finance cost for the relevant year as a percentage of Average Total Borrowings.",
    "Spread": "Spread represents Yield on Average Loans (AUM) less Average Cost of Borrowings.",
    "Net Interest Margin": "Net Interest Margin represents Net Interest Income for the relevant year as a percentage of Average Loans (AUM).",
    "Cost to Income Ratio": "Cost to Income Ratio represents operating expenses as a percentage of total income for the relevant year.",
    "CRAR": "CRAR (Capital to Risk Weighted Assets Ratio) represents the Capital Adequacy Ratio computed as per applicable RBI guidelines.",
    "Employee Attrition Rate": "Employee Attrition Rate represents the percentage of employees who exited during the relevant year relative to average headcount.",
    "Revenue per Employee": "Revenue per Employee is calculated as revenue from operations divided by the average number of employees during the relevant year.",
    "Client Retention Rate": "Client Retention Rate represents the percentage of clients retained from the preceding year.",
}


def generate_draft_docx(session: Dict[str, Any], schema: Dict[str, Any], output_path: str):
    """Generates the Draft Abridged Prospectus matching the exact cover layout and structure."""
    form_data = session.get("form_data", {})
    extracted_data = session.get("extracted_data", {})

    merged = {}
    for doc_type, data in extracted_data.items():
        if isinstance(data, dict):
            merged.update(data)
    merged.update(form_data)

    def g(key, placeholder="[●]"):
        val = merged.get(key)
        return val if not _is_empty_value(val) else placeholder

    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # Global Font Setup
    normal_font = doc.styles['Normal'].font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(9.5)
    normal_font.color.rgb = RGBColor(0, 0, 0)

    company_name = g("company_name", "[Company Name]")
    promoters = merged.get("promoters")
    promoter_names_list = merged.get("promoter_names")
    has_promoter = (isinstance(promoters, list) and len(promoters) > 0) or (isinstance(promoter_names_list, list) and len(promoter_names_list) > 0)

    # ═══════════════════════════════ PAGE 1 ═══════════════════════════════
    comp_p = doc.add_paragraph()
    comp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_comp = comp_p.add_run(str(company_name).upper())
    run_comp.font.size = Pt(12)
    run_comp.font.bold = True

    if merged.get("former_name"):
        former_p = doc.add_paragraph()
        former_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_former = former_p.add_run(f"(Formerly {merged.get('former_name')})")
        run_former.font.size = Pt(9.5)
        run_former.font.bold = True

    cin_p = doc.add_paragraph()
    cin_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cin = cin_p.add_run(f"CORPORATE IDENTITY NUMBER: {g('cin')}")
    run_cin.font.bold = True
    run_cin.font.size = Pt(10)

    _add_grid_table(doc, ["REGISTERED OFFICE", "CONTACT PERSON", "EMAIL AND TELEPHONE", "WEBSITE"], [[
        g("registered_office"),
        f"{g('company_secretary_name')}\nCompany Secretary and Compliance Officer",
        f"E-mail:\n{g('contact_email')}\nTelephone:\n{g('contact_phone')}",
        g("company_website"),
    ]], header_bg=TABLE_HEADER_TAN)

    if has_promoter:
        names_source = promoters if isinstance(promoters, list) and promoters else promoter_names_list
        names_str = ", ".join(
            (p.get("name") if isinstance(p, dict) else str(p)) for p in names_source if p
        ) if isinstance(names_source, list) else ""
        promoter_banner_text = f"OUR PROMOTERS: {names_str.upper() if names_str else '[●]'}"
    else:
        promoter_banner_text = "OUR COMPANY IS A PROFESSIONALLY MANAGED COMPANY AND DOES NOT HAVE AN IDENTIFIABLE PROMOTER"
    _add_banner(doc, promoter_banner_text)

    # Details of the Offer
    _add_banner(doc, "DETAILS OF THE OFFER")
    offer_type = "Fresh Issue and Offer for Sale" if merged.get("ofs_size_cr") else "Fresh Issue"
    _add_grid_table(doc, ["TYPE", "FRESH ISSUE SIZE", "OFFER FOR SALE SIZE", "TOTAL OFFER SIZE", "ELIGIBILITY AND RESERVATION AMONG QIBs, NIIs AND RIIs"], [[
        offer_type,
        f"Up to [●] Equity Shares of face value ₹{g('face_value_per_share')} each aggregating up to ₹{g('fresh_issue_size_cr')} Crores",
        f"Up to [●] Equity Shares of face value ₹{g('face_value_per_share')} each aggregating up to ₹{g('ofs_size_cr')} Crores" if merged.get("ofs_size_cr") else "Nil",
        f"Up to [●] Equity Shares of face value ₹{g('face_value_per_share')} each aggregating up to ₹{g('issue_size')} Crores",
        "The Offer is being made pursuant to Regulation 6(1) of the SEBI ICDR Regulations. For further details, see 'Other Regulatory and Statutory Disclosures' of the Draft Red Herring Prospectus.",
    ]], header_bg=TABLE_HEADER_TAN)

    selling_shareholders = merged.get("selling_shareholders")
    if isinstance(selling_shareholders, list) and selling_shareholders:
        _add_banner(doc, "DETAILS OF THE OFFER FOR SALE BY SELLING SHAREHOLDERS")
        _add_grid_table(doc, ["NAME OF THE SELLING SHAREHOLDER", "TYPE", "NUMBER OF EQUITY SHARES OFFERED/ AMOUNT", "WEIGHTED AVERAGE COST OF ACQUISITION PER EQUITY SHARE (IN ₹)*"], [
            [r.get("name", ""), r.get("type", ""), r.get("shares_offered", ""), r.get("waca_per_share", "")]
            for r in selling_shareholders if isinstance(r, dict)
        ], header_bg=TABLE_HEADER_TAN)

        cert_p = doc.add_paragraph()
        r_cert = cert_p.add_run(f"*As certified by the Company's Chartered Accountants, by way of their certificate dated {g('waca_ca_certificate_date')}")
        r_cert.font.size = Pt(8.5)
        r_cert.font.italic = True

    def _add_section_block(title, text):
        _add_banner(doc, title)
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10)

    has_selling_shareholders = isinstance(selling_shareholders, list) and len(selling_shareholders) > 0
    responsibility_subject = "PROMOTER SELLING SHAREHOLDERS’" if (has_promoter and has_selling_shareholders) else "SELLING SHAREHOLDERS’" if has_selling_shareholders else "SELLING SHAREHOLDERS’"

    # Standard SEBI ICDR boilerplate (page-number references generalized to [●], since
    # pagination is specific to the source document and won't match a freshly generated
    # one), verbatim from draft/Master Chains N Jewels Limited - AP_p.docx — used as the
    # fallback only. Each is a schema field (risks_first_offer_text / general_risk_text /
    # company_responsibility_text / listing_text, editable in the Wizard's Cover Page tab)
    # so a banker/legal reviewer can override the exact wording; an untouched session still
    # renders correct, regulation-standard text rather than a blank/missing section.
    default_risks_first_offer_text = (
        f"This being the first public issue of Equity Shares of our Company, there has been no formal market for the Equity Shares. The face value of "
        f"each Equity Shares is ₹{g('face_value_per_share')}. The Offer Price, Floor Price and the Cap Price determined by our Company in consultation with "
        f"the book running lead manager (“BRLM”), in accordance with the SEBI ICDR Regulations and on the basis of the assessment of market demand for the "
        f"Equity Shares by way of the Book Building Process, as stated in “Basis for Offer Price” beginning on page [●] of the Draft Red Herring Prospectus, "
        f"should not be considered to be indicative of the market price of the Equity Shares after the Equity Shares are listed. No assurance can be given "
        f"regarding an active and/or sustained trading in the Equity Shares or regarding the price at which the Equity Shares will be traded after listing."
    )
    _add_section_block("RISKS IN RELATION TO THE FIRST OFFER", merged.get("risks_first_offer_text") or default_risks_first_offer_text)

    default_general_risk_text = (
        "Investments in equity and equity-related securities involve a degree of risk and investors should not invest any funds in this Offer unless "
        "they can afford to take the risk of losing their entire investment. Bidders are advised to read the risk factors carefully before taking an "
        "investment decision in this Offer. For taking an investment decision, Bidders must rely on their own examination of our Company and the Offer, "
        "including the risks involved. The Equity Shares in the Offer have neither been recommended nor approved by Securities and Exchange Board of "
        "India (“SEBI”), nor does SEBI guarantee the accuracy or adequacy of the contents of the Draft Red Herring Prospectus. Specific attention of the "
        "Bidders is invited to “Risk Factors” beginning on page [●] of the Draft Red Herring Prospectus."
    )
    _add_section_block("GENERAL RISK", merged.get("general_risk_text") or default_general_risk_text)

    default_responsibility_text = (
        "Our Company, having made all reasonable inquiries, accepts responsibility for and confirms that the Draft Red Herring Prospectus contains all "
        "information with regard to our Company and the Offer, which is material in the context of the Offer, that the information contained in the Draft "
        "Red Herring Prospectus is true and correct in all material aspects and is not misleading in any material respect, that the opinions and intentions "
        "expressed herein are honestly held and that there are no other facts, the omission of which makes the Draft Red Herring Prospectus as a whole or "
        "any of such information or the expression of any such opinions or intentions misleading in any material respect."
    )
    if has_selling_shareholders:
        default_responsibility_text += (
            " Further, the Selling Shareholder(s) accept(s) the responsibility for and confirm(s) only the statements made or confirmed in the Draft Red "
            "Herring Prospectus to the extent of information solely in relation to the Selling Shareholder(s) and the Offered Shares and assume(s) "
            "responsibility that such statements are true and correct in all material respects and are not misleading in any material respect. The Selling "
            "Shareholder(s) assume(s) no responsibility for any other statement, disclosures and undertakings in the Draft Red Herring Prospectus, "
            "including, inter alia, any of the statements, disclosure and undertakings made by or relating to our Company or our Company’s business or "
            "any other person."
        )
    _add_section_block(f"COMPANY’S AND {responsibility_subject} ABSOLUTE RESPONSIBILITY", merged.get("company_responsibility_text") or default_responsibility_text)

    default_listing_text = (
        "The Equity Shares to be offered through Red Herring Prospectus are proposed to be listed on the BSE Limited (the “BSE”) and National Stock "
        "Exchange of India Limited (the “NSE”, and together with BSE, the “Stock Exchanges”). For the purposes of this Offer, the Designated Stock "
        "Exchange shall be [●]."
    )
    _add_section_block("LISTING", merged.get("listing_text") or default_listing_text)

    _add_banner(doc, "BOOK RUNNING LEAD MANAGER")
    _add_grid_table(doc, ["LOGO AND NAME", "CONTACT PERSON", "EMAIL AND TELEPHONE"], [[g("lead_manager"), "[●]", "[●]"]], header_bg=TABLE_HEADER_TAN)

    _add_banner(doc, "REGISTRAR TO THE OFFER")
    _add_grid_table(doc, ["LOGO AND NAME", "CONTACT PERSON", "EMAIL AND TELEPHONE"], [[g("registrar"), "[●]", "[●]"]], header_bg=TABLE_HEADER_TAN)

    _add_banner(doc, "BID/OFFER PERIOD")
    _add_grid_table(doc, ["ANCHOR INVESTOR BID/ OFFER PERIOD", "BID/ OFFER OPENS ON", "BID/ OFFER CLOSES ON"], [["[●]*", "[●]", "[●]#^"]], header_bg=TABLE_HEADER_TAN)
    if merged.get("pre_ipo_placement_amount"):
        placement_p = doc.add_paragraph()
        r_p = placement_p.add_run(
            f"@ Our Company, in consultation with the BRLM, may consider a Pre-IPO Placement aggregating up to ₹{g('pre_ipo_placement_amount')} Crores "
            f"prior to filing of the Red Herring Prospectus with the RoC, not exceeding 20% of the Fresh Issue size. {merged.get('pre_ipo_placement_terms', '')}"
        )
        r_p.font.size = Pt(8.5)
        r_p.font.italic = True

    doc.add_page_break()

    # ═══════════════════════════════ PAGE 2+ ═══════════════════════════════
    h_abr = doc.add_paragraph()
    r_abr = h_abr.add_run("IN THE NATURE OF DRAFT ABRIDGED PROSPECTUS - MEMORANDUM CONTAINING SALIENT FEATURES OF THE DRAFT RED HERRING PROSPECTUS")
    r_abr.font.bold = True
    r_abr.font.size = Pt(11)

    intro_p = doc.add_paragraph(
        "The following is a general summary of certain disclosures in the Draft Red Herring Prospectus and the terms of the Offer and is not exhaustive, "
        "nor does it purport to contain a summary of all the disclosures in the Draft Red Herring Prospectus or all details relevant to prospective investors. "
        "This summary should be read in conjunction with, and is qualified in its entirety by, the more detailed information appearing elsewhere in the Draft "
        "Red Herring Prospectus, which is available at the websites of SEBI at www.sebi.gov.in, National Stock Exchange of India Limited and BSE Limited, "
        f"the Company at {g('company_website', 'the Company website')} and the BRLM at [●].\n\n"
        "References below to page numbers are to page numbers of the Draft Red Herring Prospectus. Unless otherwise specified all capitalized terms used "
        f"herein and not specifically defined bear the same meaning as ascribed to them in the Draft Red Herring Prospectus of {company_name}."
    )
    for run in intro_p.runs:
        run.font.size = Pt(9)

    section_num = 0

    def next_num():
        nonlocal section_num
        section_num += 1
        return section_num

    # 1. Primary Business
    _heading_numbered(doc, next_num(), "Summary of primary business of our Company")
    _heading_lettered(doc, "a", "Business overview - products and services")
    business_narrative = merged.get("products_services_description") or merged.get("products_services")
    if business_narrative:
        doc.add_paragraph(_clean_llm_markdown(business_narrative))
    else:
        _missing_marker(doc, "[REQUIRES INPUT: describe products and services — use AI Assist in the wizard or enter manually]")

    _heading_lettered(doc, "b", "Industries served and typical customers")
    if merged.get("industries_served"):
        doc.add_paragraph(merged.get("industries_served"))
        if merged.get("typical_customers") and merged.get("typical_customers") not in merged.get("industries_served"):
            doc.add_paragraph(f"Our typical customers include {merged.get('typical_customers')}.")
    elif merged.get("typical_customers"):
        doc.add_paragraph(f"Our typical customers include {merged.get('typical_customers')}.")
    else:
        _missing_marker(doc, "[REQUIRES INPUT: describe industries served and typical customers]")

    _heading_lettered(doc, "c", "Segment reporting and revenue contribution")
    if merged.get("segment_reporting_applicable"):
        doc.add_paragraph("Our Company reports separate operating segments under Ind AS 108.")
    elif merged.get("segment_reporting_note"):
        doc.add_paragraph(merged.get("segment_reporting_note"))
    else:
        doc.add_paragraph("Our Company operates primarily in a single reportable segment; hence no separate segment reporting is applicable under Ind AS 108.")

    _heading_lettered(doc, "d", "Key geographies served")
    doc.add_paragraph(g("key_geographies_served", "[REQUIRES INPUT]"))

    _heading_lettered(doc, "e", "Revenue concentration among top 5 customers")
    top5 = merged.get("top5_customer_revenue_table")
    if isinstance(top5, list) and top5:
        doc.add_paragraph("Set forth below are details of the Revenue from Operations attributable to our top five customers, for the Fiscals indicated:")
        _add_grid_table(doc, ["Customer concentration", "FY (latest) ₹Cr", "FY (latest) %", "FY-1 ₹Cr", "FY-1 %", "FY-2 ₹Cr", "FY-2 %"], [
            [r.get("customer_name", ""), r.get("fy1_revenue", ""), r.get("fy1_pct", ""), r.get("fy2_revenue", ""), r.get("fy2_pct", ""), r.get("fy3_revenue", ""), r.get("fy3_pct", "")]
            for r in top5 if isinstance(r, dict)
        ])
    else:
        doc.add_paragraph("Revenue concentration among the top 5 customers is not material to the Company's business, or has not yet been provided.")

    _heading_lettered(doc, "f", "Key manufacturing or other facilities")
    facilities = merged.get("manufacturing_facility_locations")
    if isinstance(facilities, list) and facilities:
        # Group by type (Manufacturing Unit / Branch Office / Other) so this reads as proper
        # prose ("We have N manufacturing units located in X. Our branch offices are located
        # in Y.") instead of a raw comma-joined dump of every row's location string.
        by_type: Dict[str, list] = {}
        for f in facilities:
            if not isinstance(f, dict) or not f.get("location"):
                continue
            by_type.setdefault((f.get("type") or "Other").strip(), []).append(f["location"])

        sentence_parts = []
        if merged.get("registered_office"):
            sentence_parts.append(f"Our Registered Office is located at {merged.get('registered_office')}.")
        for ftype, locations in by_type.items():
            # De-duplicate for the *place names* named in the sentence (e.g. two units both in
            # Mumbai should read "...in Mumbai", not "...in Mumbai, Mumbai") while still counting
            # every row towards "We have N units".
            unique_locations = list(dict.fromkeys(locations))
            if ftype.lower().startswith("manufactur"):
                unit_word = "unit" if len(locations) == 1 else "units"
                sentence_parts.append(f"We have {len(locations)} manufacturing {unit_word} located in {', '.join(unique_locations)}.")
            elif ftype.lower().startswith("branch"):
                office_word = "office is" if len(unique_locations) == 1 else "offices are"
                sentence_parts.append(f"Our branch {office_word} located in {', '.join(unique_locations)}.")
            else:
                sentence_parts.append(f"{ftype}: {', '.join(locations)}.")
        doc.add_paragraph(" ".join(sentence_parts))
    else:
        doc.add_paragraph("Our Company does not own or operate any manufacturing facilities.")

    _heading_lettered(doc, "g", "Business strengths and strategies")
    strengths = merged.get("business_strengths")
    strategies = merged.get("business_strategies")
    if isinstance(strengths, list) and strengths:
        doc.add_paragraph("Strengths:\n" + "\n".join(f"• {s['strength']}" for s in strengths if isinstance(s, dict) and s.get("strength")))
    else:
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: business strengths]")
    if isinstance(strategies, list) and strategies:
        doc.add_paragraph("Strategies:\n" + "\n".join(f"• {s['strategy']}" for s in strategies if isinstance(s, dict) and s.get("strategy")))
    else:
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: business strategies]")

    # 2. Industry Summary
    _heading_numbered(doc, next_num(), "Summary of industry in which our Company operates")
    if merged.get("industry_growth_narrative") or merged.get("industry_narrative"):
        doc.add_paragraph(merged.get("industry_growth_narrative") or merged.get("industry_narrative"))
    if merged.get("industry_market_size") or merged.get("industry_cagr"):
        source_note = f" (Source: {g('industry_report_source')})" if merged.get("industry_report_source") else ""
        doc.add_paragraph(f"Industry Market Size: {g('industry_market_size')}. Industry CAGR: {g('industry_cagr')}.{source_note}")
    if not any([merged.get("industry_growth_narrative"), merged.get("industry_narrative"), merged.get("industry_market_size"), merged.get("industry_cagr")]):
        _missing_marker(doc, "[MISSING: Upload the industry report to populate this section]")

    # 3. Promoters
    _heading_numbered(doc, next_num(), "Promoters")
    if isinstance(promoters, list) and promoters:
        _add_grid_table(doc, ["Name", "Designation", "DIN", "Associated Since", "Education", "Experience (yrs)"], [
            [p.get("name", ""), p.get("designation", ""), p.get("din", ""), p.get("date_associated_since", ""), p.get("education_qualification", ""), p.get("years_experience", "")]
            for p in promoters if isinstance(p, dict)
        ])
        for p in promoters:
            if isinstance(p, dict) and p.get("biography_narrative"):
                doc.add_paragraph(f"{p.get('name', '')}: {p['biography_narrative']}")
    else:
        doc.add_paragraph("Our Company is a professionally managed company and does not have an identifiable promoter in terms of the SEBI ICDR Regulations and the Companies Act.")

    # 4. Objects of the Offer
    _heading_numbered(doc, next_num(), "Objects of the Offer")
    use_of_proceeds = merged.get("use_of_proceeds")
    if isinstance(use_of_proceeds, list) and use_of_proceeds:
        _add_grid_table(doc, ["Sr. No.", "Particulars", "Estimated Amount (₹ Cr)"], [
            [str(i + 1), r.get("particular", ""), r.get("estimated_amount_cr", "")] for i, r in enumerate(use_of_proceeds) if isinstance(r, dict)
        ])
    else:
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: itemized use of proceeds]")
    gross_proceeds = float(merged.get("fresh_issue_size_cr") or merged.get("issue_size") or 0)
    gcp_val = float(merged.get("general_corp_amount") or merged.get("gcp_amount_cr") or 0)
    gcp_cap = gross_proceeds * 0.25
    if gross_proceeds > 0:
        within = gcp_val <= gcp_cap + 0.01
        gcp_note = doc.add_paragraph()
        r_gcp = gcp_note.add_run(
            f"General Corporate Purposes of ₹{gcp_val:.2f} Crores {'is within' if within else 'EXCEEDS'} the SEBI ICDR Reg 230(2) cap of 25% "
            f"of Gross Proceeds (₹{gcp_cap:.2f} Crores)."
        )
        r_gcp.font.size = Pt(9)
        if not within:
            r_gcp.font.bold = True
            r_gcp.font.color.rgb = RGBColor(220, 38, 38)

    # 5. Shareholding Pattern
    _heading_numbered(doc, next_num(), "Aggregate Pre-Offer shareholding of our Promoters, Promoter Group and top 10 shareholders")
    pre_offer = merged.get("pre_offer_shareholding")
    if isinstance(pre_offer, list) and pre_offer:
        _add_grid_table(doc, ["Sr. No.", "Name of the Shareholder", "No. of Equity Shares", "% of Pre-Offer Capital"], [
            [str(i + 1), r.get("shareholder", ""), r.get("shares", ""), r.get("pct", "")] for i, r in enumerate(pre_offer) if isinstance(r, dict)
        ])
    else:
        _missing_marker(doc, "[MISSING: Upload Register of Members / cap table to populate this table]")
    promoter_group = merged.get("promoter_group_members")
    if isinstance(promoter_group, list) and promoter_group:
        doc.add_paragraph("Promoter Group Members: " + ", ".join(m.get("name", "") for m in promoter_group if isinstance(m, dict)))
    doc.add_paragraph("Post-Offer shareholding will be updated in the Prospectus once the Offer Price and Basis of Allotment are finalized.")

    # 6. Financial Information
    _heading_numbered(doc, next_num(), "Summary of Restated Financial Information")
    _add_multi_year_table(doc, [
        ("Equity Share Capital", "equity_share_capital"),
        ("Net Worth", "net_worth"),
        ("Revenue from Operations", "revenue_from_operations"),
        ("EBITDA", "ebitda"),
        ("Restated Profit After Tax", "pat"),
        ("Restated Earnings per share (basic) (in ₹)", "eps_basic"),
        ("Restated Earnings per share (diluted) (in ₹)", "eps_diluted"),
        ("Return on Net Worth (in %)", "ronw_pct"),
        ("Net Asset Value per Equity Share (in ₹)", "nav_per_share"),
        ("Total Borrowings", "total_borrowings"),
        ("Net Cash from Operating Activities", "cash_flow_operating"),
        ("Net Cash from Investing Activities", "cash_flow_investing"),
        ("Net Cash from Financing Activities", "cash_flow_financing"),
    ], merged)
    doc.add_paragraph("")
    _add_notes_block(doc, [
        "Net worth means the aggregate value of paid-up share capital and all reserves created out of profits and securities premium account, and debit or "
        "credit balance of profit and loss account, after deducting the aggregate value of accumulated losses and miscellaneous expenditure not written off, "
        "as per the restated balance sheet, but does not include reserves created out of revaluation of assets, capital reserve, foreign currency translation "
        "reserve and write-back of depreciation.",
        "Revenue from operations is as per the Restated Financial Statements.",
        "EBITDA is calculated as profit before tax plus finance costs, depreciation and amortization expense, less other income.",
        "PAT is profit for the year as per the Restated Financial Statements.",
        "Basic EPS (₹) = Total Comprehensive Income ÷ Number of shares at the end of the year.",
        "Diluted EPS (₹) = Total Comprehensive Income ÷ Number of diluted shares at the end of the year.",
        "Return on Net Worth (%) = Net profit after tax, as restated, for the financial year divided by Net Worth as restated, as at the end of the financial year.",
        "Net Asset Value per Share (NAV) = Net Worth as restated, as of the end of the financial year, divided by the number of equity shares as at the end of the financial year.",
        "Total Borrowings are calculated as the sum of non-current borrowings and current borrowings.",
    ])

    # 7. Key Performance Indicators
    _heading_numbered(doc, next_num(), "Summary of Key Performance Indicators")
    if merged.get("kpi_sector"):
        doc.add_paragraph(f"KPI Sector Template: {merged.get('kpi_sector')}")
    kpi_values = merged.get("kpi_values")
    if isinstance(kpi_values, list) and kpi_values:
        _add_grid_table(doc, ["Key Performance Indicator", "Unit", "FY (latest)", "FY-1", "FY-2"], [
            [r.get("kpi_name", ""), r.get("unit", ""), r.get("fy1_value", ""), r.get("fy2_value", ""), r.get("fy3_value", "")]
            for r in kpi_values if isinstance(r, dict)
        ])
        doc.add_paragraph("")
        kpi_names_present = [r.get("kpi_name", "") for r in kpi_values if isinstance(r, dict)]
        kpi_notes = [KPI_DEFINITIONS[name] for name in kpi_names_present if name in KPI_DEFINITIONS]
        if kpi_notes:
            _add_notes_block(doc, kpi_notes)
    else:
        _missing_marker(doc, "[REQUIRES INPUT: select a KPI sector template and populate 3-year KPI values]")

    # 8. Risk Factors
    _heading_numbered(doc, next_num(), "Risk Factors")
    doc.add_paragraph("Set forth below is a summary of certain risk factors as disclosed in the Draft Red Herring Prospectus:")
    if merged.get("risk_narrative_text"):
        doc.add_paragraph(merged.get("risk_narrative_text"))
        doc.add_paragraph("")
    if merged.get("internal_risks"):
        doc.add_paragraph("Internal Risk Factors:")
        doc.add_paragraph(merged.get("internal_risks"))
    elif not merged.get("risk_narrative_text"):
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: internal risk factors]")
    if merged.get("external_risks"):
        doc.add_paragraph("External Risk Factors:")
        doc.add_paragraph(merged.get("external_risks"))
    elif not merged.get("risk_narrative_text"):
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: external risk factors]")

    # 9. WACA
    _heading_numbered(doc, next_num(), "Weighted Average Cost of Acquisition of Equity Shares")
    waca_table = merged.get("waca_table")
    if isinstance(waca_table, list) and waca_table:
        _add_grid_table(doc, ["Shareholder", "Shares Held", "WACA/Share (₹)", "Shares Acquired (1yr)", "WACA (1yr) (₹)"], [
            [r.get("shareholder", ""), r.get("shares_held", ""), r.get("waca_per_share", ""), r.get("shares_acquired_last_1yr", ""), r.get("waca_last_1yr", "")]
            for r in waca_table if isinstance(r, dict)
        ])
        doc.add_paragraph(f"As certified per Chartered Accountant certificate dated {g('waca_ca_certificate_date')}.")
    else:
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: WACA table — must reference a dated CA certificate, cannot be auto-generated]")

    # 10. Board of Directors and KMP
    _heading_numbered(doc, next_num(), "Board of Directors and Key Managerial Personnel")
    directors = merged.get("directors")
    if isinstance(directors, list) and directors:
        _add_grid_table(doc, ["Name", "DIN", "Designation", "Independent?"], [
            [d.get("name", ""), d.get("din", ""), d.get("designation", ""), ("Yes" if d.get("independent_flag") in (True, "true", "yes", "Yes") else "No")]
            for d in directors if isinstance(d, dict)
        ])
    else:
        _missing_marker(doc, "[MISSING: Upload DIR-12 / board resolutions to populate Board of Directors]")
    kmp = merged.get("kmp")
    if isinstance(kmp, list) and kmp:
        doc.add_paragraph("Key Managerial Personnel:")
        _add_grid_table(doc, ["Name", "Designation"], [[k.get("name", ""), k.get("designation", "")] for k in kmp if isinstance(k, dict)])

    # 11. Auditor Qualifications
    _heading_numbered(doc, next_num(), "Auditor Qualifications")
    doc.add_paragraph(str(g("auditor_qualifications", "There have been no reservations, qualifications and adverse remarks in the Restated Financial Information.")))

    # 12. Litigation Summary
    _heading_numbered(doc, next_num(), "Summary Table of Outstanding Litigation")
    litigation_summary = merged.get("litigation_summary")
    if isinstance(litigation_summary, list) and litigation_summary:
        _add_grid_table(doc, ["Entity Type", "Criminal", "Tax", "Statutory/Reg.", "Civil", "Aggregate ₹Cr"], [
            [r.get("entity_type", ""), r.get("criminal_count", ""), r.get("tax_count", ""), r.get("statutory_regulatory_count", ""), r.get("civil_litigation_count", ""), r.get("aggregate_amount_cr", "")]
            for r in litigation_summary if isinstance(r, dict)
        ])
    else:
        _missing_marker(doc, "[REQUIRES BANKER/LEGAL INPUT: litigation summary — must originate from a structured legal-counsel schedule, not free-text scraping]")

    legend_paragraphs = [
        "The Equity Shares offered in the Offer have not been and will not be registered under the U.S. Securities Act or any other applicable law "
        "of the United States and, unless so registered, may not be offered or sold within the United States, except pursuant to an exemption from, "
        "or in a transaction not subject to, the registration requirements of the U.S. Securities Act and applicable U.S. state securities laws. "
        "Accordingly, the Equity Shares are being offered and sold outside the United States in “offshore transactions” as defined in and in "
        "reliance on Regulation S under the U.S. Securities Act and the applicable laws of each jurisdiction where such offers and sales are made.",
        "The Equity Shares have not been and will not be registered, listed or otherwise qualified in any other jurisdiction outside India and may "
        "not be offered or sold, and Bids may not be made by persons in any such jurisdiction, except in compliance with the applicable laws of such "
        "jurisdiction. There will be no public offering in the United States.",
        "The above information is given for the benefit of the Bidders. Our Company and the BRLM are not liable for any amendments or modification "
        "or changes in applicable laws or regulations, which may occur after the date of this Draft Abridged Prospectus. Bidders are advised to make "
        "their independent investigations and ensure that the number of Equity Shares Bid for do not exceed the applicable limits under laws or regulations.",
    ]
    for text in legend_paragraphs:
        legend_p = doc.add_paragraph(text)
        for run in legend_p.runs:
            run.font.bold = True
            run.font.size = Pt(8.5)

    # Save output
    doc.save(output_path)
    logger.info(f"Draft Abridged Prospectus generated and saved to {output_path}")